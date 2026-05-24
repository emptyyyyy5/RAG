import os
import win32com.client
import pythoncom
from docx import Document as DocxDocument
from typing import List
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import DashScopeEmbeddings
from langchain_community.vectorstores import Milvus
from PIL import Image
import easyocr
import fitz

# ==================== 配置区 ====================
DASHSCOPE_API_KEY = "sk-4e9b928c5d1848ed808e27565468f3ae"
DOCS_FOLDER = r"C:\Users\q1948\Desktop\project\rag\RAG"
OCR_FOLDER = "ImgorPDF"
MILVUS_HOST = "localhost"
MILVUS_PORT = "19530"
COLLECTION_NAME = "finance_regulations_qwen"
# ===============================================

os.environ["DASHSCOPE_API_KEY"] = DASHSCOPE_API_KEY
reader = easyocr.Reader(['ch_sim', 'en'])


# ---------- OCR 解析（图片/PDF） ----------
def ocr_image_to_text(image_path: str) -> str:
    try:
        results = reader.readtext(image_path, detail=0)
        return "\n".join(results).strip()
    except Exception as e:
        print(f"EasyOCR 识别失败: {e}")
        return ""


def pdf_to_text(pdf_path: str) -> str:
    try:
        doc = fitz.open(pdf_path)
    except Exception as e:
        print(f"打开 PDF 失败: {e}")
        return ""
    all_text = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        pix = page.get_pixmap(dpi=200)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        temp_path = f"_temp_page_{page_num}.png"
        img.save(temp_path)
        text = ocr_image_to_text(temp_path)
        if text:
            all_text.append(text)
        os.remove(temp_path)
    doc.close()
    return "\n".join(all_text)


# ---------- Word/TXT 解析 ----------
def parse_word_with_win32(file_path: str) -> str:
    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(file_path)
        text = doc.Content.Text
        return text
    except Exception as e:
        raise RuntimeError(f"Word 解析失败：{file_path}，错误：{e}")
    finally:
        if doc is not None:
            try:
                doc.Close()
            except:
                pass
        if word is not None:
            try:
                word.Quit()
            except:
                pass
        pythoncom.CoUninitialize()


def parse_text_file(file_path: str) -> str:
    for enc in ['utf-8', 'gbk', 'gb2312', 'ansi']:
        try:
            with open(file_path, 'r', encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise RuntimeError(f"无法识别文件编码：{file_path}")


# ---------- 统一解析入口 ----------
def parse_file(file_path: str) -> Document:
    ext = os.path.splitext(file_path)[-1].lower()

    if ext == '.docx':
        doc = DocxDocument(file_path)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    elif ext == '.doc':
        text = parse_word_with_win32(file_path)
    elif ext == '.txt':
        text = parse_text_file(file_path)
    elif ext == '.pdf':
        text = pdf_to_text(file_path)
    elif ext in ('.png', '.jpg', '.jpeg', '.bmp'):
        text = ocr_image_to_text(file_path)
    else:
        raise ValueError(f"不支持的文件类型：{ext}")

    if not text.strip():
        raise ValueError(f"文件内容为空：{file_path}")

    print(f"   内容长度：{len(text)} 字符")
    return Document(
        page_content=text,
        metadata={
            "source": file_path,
            "file_name": os.path.basename(file_path),
            "file_type": ext.replace('.', ''),
            "char_count": len(text)
        }
    )


def load_documents_from_folders(folder_paths: List[str]) -> List[Document]:
    supported_ext = ('.doc', '.docx', '.txt', '.pdf', '.png', '.jpg', '.jpeg', '.bmp')
    all_docs = []
    for folder_path in folder_paths:
        if not os.path.isdir(folder_path):
            print(f"⚠️ 目录不存在: {folder_path}")
            continue
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                if file.lower().endswith(supported_ext):
                    file_path = os.path.join(root, file)
                    try:
                        all_docs.append(parse_file(file_path))
                        print(f"✔ 已解析：{file_path}")
                    except Exception as e:
                        print(f"✘ 解析失败：{file_path}，错误：{e}")
    return all_docs


# ---------- 主流程 ----------
if __name__ == "__main__":
    print("=" * 50)
    print("开始使用通义千问 Embedding 重建向量库")
    print("=" * 50)

    print("\n[1/4] 加载文档...")
    docs = load_documents_from_folders([DOCS_FOLDER, OCR_FOLDER])
    print(f"共加载 {len(docs)} 个文档")

    if len(docs) == 0:
        print("没有找到任何文档，请检查文件夹路径。")
        exit()

    print("\n[2/4] 文本分块...")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=150,
        separators=[
            "\n第", "\n一、", "\n二、", "\n三、",
            "\n（一）", "\n（二）",
            "\n\n", "\n", "。", "；", " ", ""
        ]
    )
    chunks = text_splitter.split_documents(docs)
    print(f"共生成 {len(chunks)} 个文本块")

    print("\n[3/4] 初始化通义千问 Embedding 模型...")
    embeddings = DashScopeEmbeddings(model="text-embedding-v2")

    print("\n[4/4] 向量化并存入 Milvus...")
    vector_store = Milvus.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name=COLLECTION_NAME,
        connection_args={"host": MILVUS_HOST, "port": MILVUS_PORT},
        index_params={
            "index_type": "IVF_FLAT",
            "metric_type": "IP",
            "params": {"nlist": 128}
        },
        drop_old=True
    )
    print("\n✅ 向量库重建完成！")
    print(f"Collection 名称：{COLLECTION_NAME}")
