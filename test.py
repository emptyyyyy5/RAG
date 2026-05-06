# test.py 顶部增加导入
import os
import win32com.client
import pythoncom
from docx import Document as DocxDocument
from typing import List
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings  # 新版导入
from langchain_community.vectorstores import Milvus

os.environ["HF_ENDPOINT"] = "https://hf-mirror.com"
# ---------- 改进后的 Word 解析函数 ----------
def parse_word_with_win32(file_path: str) -> str:
    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(file_path)
        text = doc.Content.Text
        return text
    except Exception as e:
        raise RuntimeError(f"Word 解析失败：{file_path}，错误：{e}")
    finally:
        if doc is not None:
            try:
                doc.Close()
            except:
                pass
        if word is not None:
            try:
                word.Quit()
            except:
                pass
        pythoncom.CoUninitialize()

# ---------- 解析文件分发 ----------
def parse_file(file_path: str) -> Document:
    ext = os.path.splitext(file_path)[-1].lower()
    if ext == '.docx':
        doc = DocxDocument(file_path)
        text = '\n'.join([p.text for p in doc.paragraphs])
    elif ext == '.doc':
        text = parse_word_with_win32(file_path)
    elif ext == '.txt':
        for enc in ['utf-8', 'gbk', 'gb2312', 'ansi']:
            try:
                with open(file_path, 'r', encoding=enc) as f:
                    text = f.read()
                break
            except UnicodeDecodeError:
                continue
        else:
            raise RuntimeError(f"无法识别文件编码：{file_path}")
    else:
        raise ValueError(f"不支持的文件类型：{ext}")

    return Document(
        page_content=text,
        metadata={"source": file_path, "file_name": os.path.basename(file_path)}
    )

# ---------- 其余函数保持不变 ----------
def load_documents_from_folder(folder_path: str) -> List[Document]:
    all_docs = []
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith(('.doc', '.docx', '.txt')):
                file_path = os.path.join(root, file)
                try:
                    all_docs.append(parse_file(file_path))
                    print(f"✔ 已解析：{file_path}")
                except Exception as e:
                    print(f"✘ 解析失败：{file_path}，错误：{e}")
    return all_docs

# ---------- 主流程 ----------
if __name__ == "__main__":
    DOCS_FOLDER = r"D:\desktop\金融监督管理局\test"   # 保持你的路径

    print("开始加载文档...")
    docs = load_documents_from_folder(DOCS_FOLDER)
    print(f"共加载 {len(docs)} 个文档")

    if len(docs) == 0:
        print("没有找到任何文档，请检查文件夹路径。")
        exit()

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
        separators=["\n\n", "\n", "。", "；", "，", " ", ""]
    )
    chunks = text_splitter.split_documents(docs)
    print(f"共生成 {len(chunks)} 个文本块")

    print("正在加载本地 Embedding 模型...")
    embeddings = HuggingFaceEmbeddings(
        model_name="D:/desktop/pycharm/bge-small-zh-v1.5",
        model_kwargs={'device': 'cpu'},
        encode_kwargs={'normalize_embeddings': True}
    )

    print("正在连接 Milvus 并创建向量库...")
    vector_store = Milvus.from_documents(
        documents=chunks,
        embedding=embeddings,
        collection_name="finance_regulations",
        connection_args={"host": "localhost", "port": "19530"},
        index_params={
            "index_type": "IVF_FLAT",
            "metric_type": "IP",
            "params": {"nlist": 1024}
        },
        drop_old=False
    )
    print("✅ 知识库构建完成！")